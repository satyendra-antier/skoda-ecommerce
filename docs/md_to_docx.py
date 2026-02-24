#!/usr/bin/env python3
"""Convert the one-page report markdown to .docx for Google Docs."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_table_from_rows(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
        t.rows[0].cells[j].paragraphs[0].runs[0].bold = True
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < len(t.rows[i + 1].cells):
                t.rows[i + 1].cells[j].text = str(cell)
    return t

def main():
    doc = Document()
    doc.add_paragraph()
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('ŠKODA Lifestyle — E-Commerce Integration · One-Page Report')
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    sub = doc.add_paragraph('Phase 1 · Version 1.0 · Confidential')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph('—' * 40)
    doc.add_paragraph()

    add_heading(doc, 'What We Need to Build', level=1)
    add_table_from_rows(doc, ['Area', 'Scope'], [
        ('Front-end', 'Product listing (grid, image, name, price, stock status), product detail (gallery, specs, quantity, Add to Cart), cart (summary, quantity controls, totals, checkout CTA), checkout (Full Name, Mobile, Email, Shipping Address, City, State, Pincode — with validation). Design reference: Porsche Shop (https://shop.porsche.com/us/en-US).'),
        ('Payments', 'BillDesk integration (server-side): unique Order ID, secure hash (server-only), redirect to BillDesk hosted page, callback handling, response validation, order status update. No card/sensitive data stored. Test + Production.'),
        ('Inventory', 'Centralised model: Product ID, SKU, Price, Available Stock, Status (Active/Out of Stock). Deduct stock only after successful payment; oversell protection; out-of-stock = no Add to Cart, no checkout.'),
        ('CRM', 'Zoho: auto-create/update order on successful payment. Fields: Customer Name, Mobile, Email, Shipping Address, Products, Quantity, Order ID, Amount Paid, Payment Status, Order Date/Time. Deduplication, error logging, optional retry.'),
        ('Admin', 'Web dashboard: view/filter orders, update stock, override status, CSV export, payment status summary.'),
        ('Security', 'Hash server-side only; credentials in env; SSL on payment pages; audit logs for order/payment; no card storage.'),
    ])
    doc.add_paragraph()

    add_heading(doc, 'End-to-End Flow', level=1)
    flow = """Product listing → Product detail → Add to Cart → Cart → Checkout (customer + address)
       → Stock check → BillDesk redirect → Payment (hosted) → Callback
       → Validate response → Update order status → Deduct inventory
       → Create/update Zoho order → Confirmation to customer"""
    p = doc.add_paragraph(flow)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Consolas'
    add_para(doc, 'Payment statuses: Pending → Successful / Failed (order and UI updated accordingly).', bold=True)
    add_para(doc, 'Out-of-stock: Product hidden/disabled for Add to Cart; checkout blocked if stock depletes between cart and payment.', bold=True)
    doc.add_paragraph()

    add_heading(doc, 'Assumptions', level=1)
    for bullet in [
        'Client is already onboarded with BillDesk; Test and Production credentials will be provided.',
        'Single centralised inventory (no dealer/multi-warehouse in Phase 1).',
        'Zoho module name, field mapping, and API credentials will be provided at kick-off.',
        'Design and UX will follow Porsche Shop as reference; no separate design phase specified.',
        'Retry for failed Zoho sync is "if feasible within scope."',
        'Server environment, deployment access, SSL, and domain are client-provided.',
    ]:
        doc.add_paragraph(bullet, style='List Bullet')
    doc.add_paragraph()

    add_heading(doc, 'Out of Scope (Phase 1)', level=1)
    add_para(doc, 'Dealer-level inventory · ERP integration · Multi-warehouse · Automated refunds · Advanced discount/coupon engine.')
    doc.add_paragraph()

    add_heading(doc, 'Client Inputs Required at Kick-off', level=1)
    add_table_from_rows(doc, ['From', 'Items'], [
        ('BillDesk', 'Merchant ID, Secret Key, Test credentials, Integration docs, Approved callback URL'),
        ('Zoho', 'API credentials, Order module name, Field mapping sheet'),
        ('IT', 'Server env details, Deployment access, SSL confirmation, Domain details'),
    ])
    doc.add_paragraph()

    add_heading(doc, 'Deliverables', level=1)
    add_para(doc, 'Functional e-commerce journey · BillDesk (Test & Prod) · Zoho CRM integration · Inventory module · Admin dashboard · UAT & go-live support · Basic technical documentation.')
    doc.add_paragraph()
    doc.add_paragraph('—' * 40)
    doc.add_paragraph()
    footer = doc.add_paragraph('ŠKODA AUTO India · ŠKODA Lifestyle E-Commerce Integration · Phase 1 SOW')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.italic = True
        run.font.size = Pt(9)

    out = '/home/user/Desktop/skoda/SKODA-LIFESTYLE-ONE-PAGE-REPORT.docx'
    doc.save(out)
    print('Saved:', out)

if __name__ == '__main__':
    main()
